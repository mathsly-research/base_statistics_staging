# -*- coding: utf-8 -*-
import streamlit as st
import pandas as pd
import numpy as np
import json, io, uuid, datetime as dt

from core.state import init_state

# Plotly
import plotly.graph_objects as go

# Opzionali per export HTML/DOCX
try:
    import markdown as mdconv  # pip install markdown
    _HAS_MD = True
except Exception:
    _HAS_MD = False

try:
    from docx import Document  # pip install python-docx
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False


# ===========================================================
# Utilità
# ===========================================================
def _ensure_report_store():
    if "report_items" not in st.session_state or not isinstance(st.session_state.report_items, list):
        st.session_state.report_items = []
    # backfill id/timestamp per elementi già presenti
    changed = False
    for it in st.session_state.report_items:
        if "id" not in it:
            it["id"] = str(uuid.uuid4())
            changed = True
        if "created_at" not in it:
            it["created_at"] = dt.datetime.now().isoformat(timespec="seconds")
            changed = True
        if "title" not in it:
            it["title"] = it.get("type","Item")
            changed = True
        if "content" not in it:
            it["content"] = {}
            changed = True
    if changed:
        st.session_state.report_items = list(st.session_state.report_items)

def _to_table(obj, round_at=4):
    """Converte dict/list in DataFrame leggibile (best-effort)."""
    if obj is None:
        return pd.DataFrame()
    try:
        if isinstance(obj, pd.DataFrame):
            return obj
        if isinstance(obj, list):
            if len(obj)==0:
                return pd.DataFrame()
            if isinstance(obj[0], dict):
                return pd.DataFrame(obj)
            return pd.DataFrame({ "valori": obj })
        if isinstance(obj, dict):
            return pd.DataFrame(list(obj.items()), columns=["Chiave","Valore"])
    except Exception:
        pass
    return pd.DataFrame({"Valore":[obj]})

def _short_stats_for_card(item):
    """Estratto breve per la card, dipendente dal tipo."""
    t = item.get("type","").lower()
    c = item.get("content",{})
    out = []
    try:
        if t == "regression_ols":
            out.append(f"R²={c.get('r2',np.nan):.3f} | AIC={c.get('aic',np.nan):.1f}")
        elif t == "regression_logit":
            pr2 = c.get("pseudo_r2_mcfadden", None)
            if pr2 is not None:
                out.append(f"Pseudo-R²={pr2:.3f}")
            aic = c.get("aic", None)
            if aic is not None and np.isfinite(aic):
                out.append(f"AIC={aic:.1f}")
        elif t == "regression_poisson":
            out.append(f"AIC={c.get('aic',np.nan):.1f} | Dev={c.get('deviance',np.nan):.1f}")
        elif t == "regression_logit_regularized":
            out.append(f"Penalità={c.get('penalty','?')} | C={c.get('C','?')}")
        elif t == "diagnostic_test":
            cm = c.get("confusion_matrix",{})
            prev = cm.get("Prevalence", None)
            if prev is not None and np.isfinite(prev):
                out.append(f"Prev={prev:.2%}")
            try:
                mets = pd.DataFrame(c.get("metrics",[]))
                for m in ["Sensibilità","Specificità","PPV","NPV"]:
                    val = float(mets.loc[mets["Metrica"]==m, "Valore"].values[0])
                    out.append(f"{m[:3]}={val:.2f}")
                    if len(out)>=3: break
            except Exception:
                pass
        elif t == "agreement_continuous":
            ba = pd.DataFrame(item["content"].get("bland_altman",[]))
            if not ba.empty:
                try:
                    bias = float(ba.loc[ba["Parametro"]=="Bias","Valore"].values[0])
                    loa_l = float(ba.loc[ba["Parametro"]=="LoA bassa","Valore"].values[0])
                    loa_h = float(ba.loc[ba["Parametro"]=="LoA alta","Valore"].values[0])
                    out.append(f"Bias={bias:.3g} | LoA=[{loa_l:.3g},{loa_h:.3g}]")
                except Exception:
                    pass
            ccc = item["content"].get("ccc",{})
            if "value" in ccc:
                out.append(f"CCC={ccc['value']:.3f}")
        elif t == "agreement_icc":
            icc21 = item["content"].get("icc21",None)
            icc2k = item["content"].get("icc2k",None)
            out.append(f"ICC(2,1)={icc21:.3f} | ICC(2,k)={icc2k:.3f}")
        elif t == "agreement_kappa":
            out.append(f"Kappa={item['content'].get('kappa',np.nan):.3f}")
        elif t == "survival_analysis":
            N = item["content"].get("N",None); ev = item["content"].get("events",None)
            if N is not None and ev is not None:
                out.append(f"N={N}, Eventi={ev}")
        elif t == "longitudinal_lmm":
            r2m = item["content"].get("r2_marginal", None)
            r2c = item["content"].get("r2_conditional", None)
            out.append(f"R²m={r2m:.3f} | R²c={r2c:.3f}")
        elif t == "longitudinal_gee":
            out.append(f"Corr={item['content'].get('cov_struct','?')}")
    except Exception:
        pass
    return " • ".join([s for s in out if s])

def _render_small_bar_by_type(df_counts):
    fig = go.Figure()
    fig.add_trace(go.Bar(x=df_counts["type"], y=df_counts["n"]))
    fig.update_layout(title="Elementi per tipologia", xaxis_title="", yaxis_title="Conteggio", height=300, margin=dict(t=40,b=20))
    return fig

def _compile_markdown(title, author, selected_items):
    now = dt.datetime.now().strftime("%Y-%m-%d %H:%M")
    md = [f"# {title}", "", f"*Generato: {now}*  ", ""]
    if author:
        md.insert(2, f"*Autore: {author}*  ")

    # Sezione dataset (se presente)
    if "df_original" in st.session_state and isinstance(st.session_state.df_original, pd.DataFrame):
        df = st.session_state.df_original
        md += [
            "## Dataset",
            f"- Osservazioni: **{df.shape[0]}**  ",
            f"- Variabili: **{df.shape[1]}**  ",
            ""
        ]

    for it in selected_items:
        md += [f"## {it.get('title','Sezione')}", f"*Tipo:* `{it.get('type','-')}`  ", f"*Creato:* {it.get('created_at','-')}  ", ""]
        c = it.get("content",{})
        # prova formattazioni sintetiche per tipi noti
        t = it.get("type","").lower()
        try:
            if t == "regression_ols":
                md += ["**OLS**",
                       f"- N: {c.get('nobs','')}",
                       f"- R²: {c.get('r2',''):.3f} (adj {c.get('r2_adj',''):.3f})",
                       f"- AIC: {c.get('aic',''):.2f} | BIC: {c.get('bic',''):.2f}", ""]
                coefs = _to_table(c.get("coefficients",[])).round(4)
                md += ["**Coefficienti**", coefs.to_markdown(index=False), ""]
            elif t == "regression_logit":
                md += ["**Logistica**",
                       f"- N: {c.get('nobs','')}",
                       f"- LogLik: {c.get('loglik','')} | Null: {c.get('loglik_null','')}",
                       f"- Pseudo-R² (McFadden): {c.get('pseudo_r2_mcfadden',np.nan):.3f}",
                       f"- AIC: {c.get('aic',''):.2f} | BIC: {c.get('bic',''):.2f}", ""]
                ors = _to_table(c.get("odds_ratios",[])).round(4)
                md += ["**Odds Ratio**", ors.to_markdown(index=False), ""]
            elif t == "regression_logit_regularized":
                md += [f"**Logistica regolarizzata** (penalty={c.get('penalty')}, C={c.get('C')})",""]
                coef = _to_table(c.get("coefficients",[])).round(4)
                md += ["**Coefficienti**", coef.to_markdown(index=False), ""]
            elif t == "regression_poisson":
                md += [f"**GLM Poisson** — N={c.get('nobs','')}  ",
                       f"- Deviance: {c.get('deviance','')} | Pearson χ²: {c.get('pearson_chi2','')} | AIC: {c.get('aic','')}", ""]
                irr = _to_table(c.get("irr",[])).round(4)
                md += ["**IRR**", irr.to_markdown(index=False), ""]
            elif t == "diagnostic_test":
                cm = c.get("confusion_matrix",{})
                md += [f"**Test diagnostico** — N={cm.get('Total','')}, Prevalenza={cm.get('Prevalence',np.nan):.3f}", ""]
                mets = _to_table(c.get("metrics",[])).round(3)
                md += ["**Metriche (IC95%)**", mets.to_markdown(index=False), ""]
                if "calibration" in c:
                    cal = c["calibration"]
                    md += [f"**Calibrazione** — Brier={cal.get('brier',np.nan):.3f}, Intercetta={cal.get('intercept',np.nan):.3f}, Slope={cal.get('slope',np.nan):.3f}",
                           ""]
            elif t == "agreement_continuous":
                ba = _to_table(c.get("bland_altman",[])).round(4)
                md += ["**Bland–Altman**", ba.to_markdown(index=False), ""]
                ccc = c.get("ccc",{})
                if ccc:
                    md += [f"**CCC (Lin)**: {ccc.get('value',np.nan):.3f}  (IC 95%: {ccc.get('ci',[np.nan,np.nan])[0]:.3f}–{ccc.get('ci',[np.nan,np.nan])[1]:.3f})", ""]
            elif t == "agreement_icc":
                md += [f"**ICC** — ICC(2,1)={c.get('icc21',np.nan):.3f}, ICC(2,k)={c.get('icc2k',np.nan):.3f}", ""]
            elif t == "agreement_kappa":
                md += [f"**Kappa** — {c.get('kappa',np.nan):.3f} (pesata: {c.get('weighted',False)})", ""]
            elif t == "survival_analysis":
                md += [f"**Sopravvivenza** — N={c.get('N','')}, Eventi={c.get('events','')}, Censure={c.get('censored','')}", ""]
                if "cox_hr" in c:
                    hr = _to_table(c.get("cox_hr",[])).round(4)
                    md += ["**Cox — HR (IC95%)**", hr.to_markdown(index=False), ""]
            elif t == "longitudinal_lmm":
                md += [f"**LMM** — formula: `{c.get('formula','')}`; random: `{c.get('random','')}`",
                       f"- R² marginale: {c.get('r2_marginal',np.nan):.3f} | R² condizionale: {c.get('r2_conditional',np.nan):.3f}", ""]
                fe = _to_table(c.get("fixed_effects",[])).round(4)
                md += ["**Effetti fissi**", fe.to_markdown(index=False), ""]
            elif t == "longitudinal_gee":
                md += [f"**GEE** — formula: `{c.get('formula','')}`; corr: `{c.get('cov_struct','')}`",""]
                cf = _to_table(c.get("coefficients",[])).round(4)
                md += ["**Coefficienti**", cf.to_markdown(index=False), ""]
            else:
                md += ["**Dettagli**", _to_table(c).to_markdown(index=False), ""]
        except Exception:
            md += ["**Dettagli (grezzi)**", "```json", json.dumps(c, indent=2, ensure_ascii=False), "```", ""]
    return "\n".join(md)


def _build_html_from_markdown(md_text:str) -> str:
    if _HAS_MD:
        body = mdconv.markdown(md_text, extensions=['tables', 'fenced_code'])
    else:
        # fallback minimale
        body = "<pre style='white-space:pre-wrap'>" + md_text.replace("&","&amp;").replace("<","&lt;") + "</pre>"
    html = f"""<!doctype html>
<html><head>
<meta charset="utf-8"/>
<title>Report</title>
<style>
body{{font-family: -apple-system, Segoe UI, Roboto, Helvetica, Arial, sans-serif; margin: 40px;}}
table{{border-collapse: collapse;}}
th,td{{border:1px solid #ccc; padding:6px 8px;}}
h1,h2,h3{{margin-top:1.2em;}}
code, pre{{background:#f7f7f7;}}
</style>
</head><body>{body}</body></html>"""
    return html


# ===========================================================
# Pagina
# ===========================================================
init_state()
st.title("📋 Results Summary — Dashboard")

_ensure_report_store()
items = st.session_state.report_items

# CSS per pulsanti più bassi
st.markdown("""
<style>
.small-btn button[kind="primary"], .small-btn button[kind="secondary"] { padding: 0.2rem 0.6rem; line-height: 1.0; }
.card { border-radius: 12px; padding: 12px; margin-bottom: 12px; box-shadow: 1px 1px 8px rgba(0,0,0,0.06); }
.card h4 { margin: 0 0 6px 0; }
.card .meta { color:#666; font-size:0.85em; margin-bottom:8px; }
</style>
""", unsafe_allow_html=True)

# =========================
# Panoramica
# =========================
st.subheader("Panoramica")
if len(items) == 0:
    st.info("Nessun elemento nel Results Summary. Aggiunga risultati dagli altri moduli.")
else:
    dfc = pd.DataFrame([{"type": it.get("type","?")} for it in items]).value_counts().reset_index()
    if dfc.empty:
        dfc = pd.DataFrame({"type":["?"],"n":[len(items)]})
    else:
        dfc.columns = ["type","n"]
    left, right = st.columns(2)
    with left:
        with st.spinner("Genero il riepilogo per tipologia…"):
            fig = _render_small_bar_by_type(dfc)
            st.plotly_chart(fig, use_container_width=True)
        st.caption("Conteggio degli elementi salvati per ciascuna tipologia.")
    with right:
        with st.spinner("Creo una tabella di sintesi…"):
            resume = pd.DataFrame([{
                "ID": it["id"][:8],
                "Tipo": it.get("type",""),
                "Titolo": it.get("title",""),
                "Creato": it.get("created_at","")
            } for it in items])
            st.dataframe(resume, use_container_width=True, hide_index=True)
        st.caption("Elenco compatto degli elementi presenti.")

# =========================
# Filtri e ricerca
# =========================
st.subheader("Filtri")
types_avail = sorted({ it.get("type","") for it in items })
colf1, colf2 = st.columns([2,3])
with colf1:
    sel_types = st.multiselect("Tipi inclusi", options=types_avail, default=types_avail)
with colf2:
    query = st.text_input("Ricerca nel titolo", value="")

filtered = [it for it in items if it.get("type","") in sel_types and (query.lower() in it.get("title","").lower())]

st.markdown(f"**Elementi filtrati**: {len(filtered)}")

# =========================
# Elenco dettagli (cards)
# =========================
for it in filtered:
    with st.container():
        st.markdown('<div class="card">', unsafe_allow_html=True)
        c1, c2 = st.columns([3,2])
        with c1:
            st.markdown(f"### {it.get('title','(senza titolo)')}")
            st.markdown(f"<div class='meta'>Tipo: <code>{it.get('type','')}</code> • Creato: {it.get('created_at','')}</div>", unsafe_allow_html=True)
            st.write(_short_stats_for_card(it))
        with c2:
            # azioni
            new_title = st.text_input("Rinomina", value=it.get("title",""), key=f"title_{it['id']}")
            cols_btn = st.columns(3)
            with cols_btn[0]:
                if st.button("💾 Salva", key=f"save_{it['id']}", help="Aggiorna il titolo"):
                    it["title"] = new_title
                    st.success("Titolo aggiornato.")
            with cols_btn[1]:
                payload = json.dumps(it, ensure_ascii=False, indent=2)
                st.download_button("⬇️ JSON", payload, file_name=f"{it['id']}_{it.get('type','item')}.json", mime="application/json", key=f"dl_{it['id']}")
            with cols_btn[2]:
                if st.button("🗑️ Elimina", key=f"del_{it['id']}"):
                    st.session_state.report_items = [jj for jj in st.session_state.report_items if jj["id"] != it["id"]]
                    st.warning("Elemento eliminato.")
                    st.stop()
        # anteprima dettagli
        with st.expander("Anteprima dettagli (tabella)", expanded=False):
            cont = it.get("content",{})
            # prova a mostrare una tabella significativa
            t = it.get("type","").lower()
            df_show = pd.DataFrame()
            if t in ("regression_ols","regression_poisson","regression_logit"):
                df_show = _to_table(cont.get("coefficients") or cont.get("irr") or cont.get("odds_ratios"))
            elif t == "diagnostic_test":
                df_show = _to_table(cont.get("metrics"))
            elif t == "agreement_continuous":
                df_show = _to_table(cont.get("bland_altman"))
            elif t == "survival_analysis":
                df_show = _to_table(cont.get("cox_hr"))
            elif t == "longitudinal_lmm":
                df_show = _to_table(cont.get("fixed_effects"))
            elif t == "longitudinal_gee":
                df_show = _to_table(cont.get("coefficients"))
            else:
                df_show = _to_table(cont)
            if not df_show.empty:
                st.dataframe(df_show.round(4), use_container_width=True)
            else:
                st.info("Nessuna tabella sintetica disponibile per questo elemento.")
        st.markdown('</div>', unsafe_allow_html=True)

st.divider()

# =========================
# Export del Report
# =========================
st.subheader("Esporta report finale")
colE1, colE2 = st.columns([3,2])
with colE1:
    report_title = st.text_input("Titolo report", value="Risultati dell’analisi")
    author = st.text_input("Autore (opzionale)", value="")
with colE2:
    include_ids = st.multiselect("Seleziona elementi da includere", options=[f"{it['id'][:8]} — {it.get('title','')}" for it in filtered], default=[f"{it['id'][:8]} — {it.get('title','')}" for it in filtered])
    chosen = []
    chosen_ids = {x.split(" — ")[0] for x in include_ids}
    for it in filtered:
        if it["id"][:8] in chosen_ids:
            chosen.append(it)

with st.spinner("Compilo il report (Markdown)…"):
    md_text = _compile_markdown(report_title, author, chosen)

# download: Markdown
st.download_button("⬇️ Scarica Markdown", md_text, file_name="report.md", mime="text/markdown")

# download: HTML
html_text = _build_html_from_markdown(md_text)
st.download_button("⬇️ Scarica HTML", html_text, file_name="report.html", mime="text/html")

# download: DOCX (se disponibile)
if _HAS_DOCX:
    with st.spinner("Genero DOCX…"):
        try:
            doc = Document()
            for line in md_text.split("\n"):
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("**") and line.endswith("**") and len(line) < 100:
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip("*"))
                    run.bold = True
                elif line.strip().startswith("|") and line.strip().endswith("|"):
                    # tabella markdown semplice
                    rows = [r.strip() for r in line.strip().split("\n") if r.strip()]
                    # qui semplifichiamo: rimandiamo la tabella, poiché la conversione completa è onerosa
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph(line)
            bio = io.BytesIO()
            doc.save(bio)
            st.download_button("⬇️ Scarica DOCX", data=bio.getvalue(), file_name="report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception as e:
            st.warning(f"DOCX non generato: {e}")
else:
    st.caption("Per l’esportazione DOCX installare **python-docx** (`pip install python-docx`).")

# =========================
# Pulisci tutto
# =========================
with st.expander("Operazioni di manutenzione", expanded=False):
    c1, c2 = st.columns(2)
    with c1:
        if st.button("🧹 Svuota Results Summary"):
            st.session_state.report_items = []
            st.success("Archivio svuotato.")
            st.stop()
    with c2:
        if st.button("💾 Esporta archivio (.json)"):
            payload = json.dumps(st.session_state.report_items, ensure_ascii=False, indent=2)
            st.download_button("⬇️ Scarica archivio", data=payload, file_name="results_summary_archive.json", mime="application/json", key="dl_archive")

# =========================
# Spiegazione
# =========================
with st.expander("ℹ️ Come leggere e usare la dashboard", expanded=False):
    st.markdown("""
**Cosa vede**  
- **Panoramica**: conteggio per tipologia.  
- **Filtri**: chiuda/mostri elementi per **tipo** e per testo nel **titolo**.  
- **Cards**: per ciascun elemento, un **riassunto** (es. R², AIC, HR, LR, ICC…) e un’**anteprima** tabellare.  
- **Azioni**: *Rinomina*, *Elimina*, *Scarica JSON* dell’elemento.  

**Export del report**  
- Scelga **Titolo**/**Autore**, selezioni gli elementi da includere, quindi scarichi il **Markdown** o l’**HTML** (DOCX se disponibile).  
- Il report è **autocontenuto** e include tabelle chiave (coefficienti, HR, metriche diagnostiche, Bland–Altman, ecc.).  

**Suggerimenti**  
- Mantenga i titoli brevi ma informativi (es. “Logit — outcome Y (positiva=1)”).  
- Se necessario, esporti l’**archivio JSON** per un backup; può ricaricarlo in futuro (basta leggere il file e reinserire in `st.session_state.report_items`).  
""")
